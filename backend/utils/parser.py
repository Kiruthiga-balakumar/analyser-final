"""
Resume and document parsing utilities.
Handles PDF, DOCX, DOC, and plain text files.
Universal & FastAPI compatible.
"""

import logging
import re
import tempfile
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger(__name__)

# =================================================
# CORE EXTRACTORS (PATH BASED)
# =================================================

def extract_text_from_pdf(file_path: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber not installed. pip install pdfplumber")

    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                else:
                    logger.warning(f"Empty text on PDF page {i+1}")
    except Exception as e:
        logger.error(f"PDF read error: {e}")
        raise
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. pip install python-docx")

    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        logger.error(f"DOCX read error: {e}")
        raise
    return text.strip()


def extract_text_from_doc(file_path: str) -> str:
    logger.warning("DOC format detected, attempting DOCX fallback")
    return extract_text_from_docx(file_path)


def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="latin-1") as f:
            return f.read().strip()


# =================================================
# UNIVERSAL ENTRY (UPLOAD SAFE)
# =================================================

def extract_text_from_upload(filename: str, file_bytes: bytes) -> str:
    """
    Extract text from uploaded file bytes (FastAPI / Streamlit safe).
    """
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file_bytes)
        temp_path = tmp.name

    try:
        return extract_text_from_file(temp_path)  # use universal wrapper
    finally:
        try:
            Path(temp_path).unlink()
        except Exception as e:
            logger.warning(f"Temp file cleanup failed: {e}")


# =================================================
# UNIVERSAL WRAPPER (FOR BACKEND IMPORT)
# =================================================

def extract_text_from_file(file_path: str) -> str:
    """
    Automatically detect file type and extract text.
    Compatible with FastAPI routes.
    """
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        return extract_text_from_docx(file_path)
    elif suffix == ".doc":
        return extract_text_from_doc(file_path)
    elif suffix in [".txt", ".text"]:
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


# =================================================
# CLEANING & VALIDATION
# =================================================

def clean_text(text: str) -> str:
    """
    Remove extra whitespace, URLs, emails, phone numbers, and special chars.
    """
    text = " ".join(text.split())
    text = re.sub(r"http\S+|www\S+", "", text)
    text = re.sub(r"\S+@\S+", "", text)
    text = re.sub(r"\b\d{3}[-.]?\d{3}[-.]?\d{4}\b", "", text)
    text = re.sub(r"[^\w\s\-\+\#\.]", " ", text)
    return " ".join(text.split()).strip()


def validate_text(text: str, min_length: int = 50) -> Tuple[bool, Optional[str]]:
    """
    Validate text length and content.
    Returns (is_valid, error_message)
    """
    if not text:
        return False, "Extracted text is empty"

    if len(text) < min_length:
        return False, f"Text too short (min {min_length} chars)"

    alpha_ratio = sum(c.isalnum() or c.isspace() for c in text) / len(text)
    if alpha_ratio < 0.5:
        return False, "Text appears corrupted"

    return True, None
